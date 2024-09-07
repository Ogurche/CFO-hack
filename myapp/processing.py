from docx import Document
# from reportlab.pdfgen import canvas
# import pandas as pd
# import matplotlib.pyplot as plt
# from PyPDF2 import PdfReader, PdfWriter
# from reportlab.lib.pagesizes import letter
# from io import BytesIO
# from reportlab.pdfbase.ttfonts import TTFont
# from reportlab.lib import colors
# from reportlab.pdfbase import pdfmetrics
# import pdfrw

def process_doc_file(file_path) -> list:
    doc= Document(file_path)
    res = []
    # Получаем текст всех абзацев в документе
    # количество абзацев в документе
    for i in range(len(doc.paragraphs)):
        res.append(doc.paragraphs[i].text)
    return res
